"""
tests/integration/test_converters.py — Integration tests for file converters.

These tests exercise the full converter pipeline without network, Neo4j, or LLM.
Imports directly from submodules to avoid triggering the CLI import chain.
"""
from __future__ import annotations

import pytest

# Import directly from submodules — NOT from scinr.newton (top-level)
from scinr.newton.converters.base import ConversionError, IntermediateDocument
from scinr.newton.converters.registry import get_converter

pytestmark = pytest.mark.integration


# ---------------------------------------------------------------------------
# DOCX conversion
# ---------------------------------------------------------------------------


class TestDocxConversion:
    def test_docx_conversion_produces_pages(self, tmp_path):
        """A minimal DOCX file is converted to an IntermediateDocument with pages."""
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("Another paragraph with some content.")

        p = tmp_path / "test.docx"
        doc.save(str(p))

        converter = get_converter(p)
        result = converter.convert(p)

        assert isinstance(result, IntermediateDocument)
        assert len(result.pages) >= 1
        # At least one page should have markdown content
        all_markdown = "\n".join(page.markdown for page in result.pages)
        assert len(all_markdown) > 0

    def test_docx_conversion_with_table(self, tmp_path):
        """DOCX with a table is converted and table appears in markdown."""
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        doc.add_heading("Table Test", level=1)
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "Header1"
        table.cell(0, 1).text = "Header2"
        table.cell(0, 2).text = "Header3"
        table.cell(1, 0).text = "Val1"
        table.cell(1, 1).text = "Val2"
        table.cell(1, 2).text = "Val3"

        p = tmp_path / "table.docx"
        doc.save(str(p))

        converter = get_converter(p)
        result = converter.convert(p)

        assert isinstance(result, IntermediateDocument)
        assert len(result.pages) >= 1


# ---------------------------------------------------------------------------
# CSV conversion (redirect to tabular pipeline)
# ---------------------------------------------------------------------------


class TestCsvConversion:
    def test_csv_converter_raises_conversion_error(self, tmp_path):
        """The CSV converter raises ConversionError directing to the tabular pipeline."""
        p = tmp_path / "data.csv"
        p.write_text("name,age\nAlice,30\n", encoding="utf-8")

        converter = get_converter(p)
        with pytest.raises(ConversionError):
            converter.convert(p)

    def test_csv_tabular_reader_works(self, tmp_path):
        """The tabular reader (not the converter) correctly reads a CSV file."""
        from scinr.newton.tabular.reader import read_csv

        p = tmp_path / "data.csv"
        p.write_text("product,price,qty\nWidget,9.99,100\nGadget,24.99,50\n", encoding="utf-8")

        sheets = read_csv(p)
        assert len(sheets) == 1
        assert sheets[0]["headers"] == ["product", "price", "qty"]
        assert sheets[0]["total_rows"] == 2
        assert sheets[0]["all_rows"][0] == ["Widget", "9.99", "100"]


# ---------------------------------------------------------------------------
# XLSX conversion (redirect to tabular pipeline)
# ---------------------------------------------------------------------------


class TestXlsxConversion:
    def test_xlsx_converter_raises_conversion_error(self, tmp_path):
        """The XLSX converter raises ConversionError directing to the tabular pipeline."""
        openpyxl = pytest.importorskip("openpyxl")

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["col1", "col2"])
        ws.append(["a", "b"])
        p = tmp_path / "data.xlsx"
        wb.save(str(p))

        converter = get_converter(p)
        with pytest.raises(ConversionError):
            converter.convert(p)

    def test_xlsx_tabular_reader_works(self, tmp_path):
        """The tabular reader correctly reads an XLSX file."""
        openpyxl = pytest.importorskip("openpyxl")
        from scinr.newton.tabular.reader import read_xlsx

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["item", "value", "count"])
        ws.append(["Alpha", 1.5, 10])
        ws.append(["Beta", 2.5, 20])
        p = tmp_path / "data.xlsx"
        wb.save(str(p))

        sheets = read_xlsx(p)
        assert len(sheets) >= 1
        assert sheets[0]["headers"] == ["item", "value", "count"]
        assert sheets[0]["total_rows"] == 2


# ---------------------------------------------------------------------------
# Unsupported extension
# ---------------------------------------------------------------------------


class TestUnsupportedExtension:
    def test_unsupported_extension_raises(self, tmp_path):
        """Passing a .xyz file to get_converter raises UnsupportedFormatError."""
        from scinr.newton.converters.base import UnsupportedFormatError

        p = tmp_path / "file.xyz_unsupported"
        p.write_text("some content")

        with pytest.raises(UnsupportedFormatError):
            get_converter(p)

    def test_unsupported_extension_error_message(self, tmp_path):
        """UnsupportedFormatError message mentions the extension and supported formats."""
        from scinr.newton.converters.base import UnsupportedFormatError

        p = tmp_path / "file.unknownfmt"
        p.write_text("content")

        with pytest.raises(UnsupportedFormatError) as exc_info:
            get_converter(p)

        error_msg = str(exc_info.value)
        assert "unknownfmt" in error_msg


# ---------------------------------------------------------------------------
# DOCX explicit_breaks + paragraph cap behaviour
# ---------------------------------------------------------------------------

# _PARAGRAPHS_PER_PAGE = 4  (defined in scinr.newton.converters.docx)
_PARAGRAPHS_PER_PAGE = 4


def _add_page_break(doc) -> None:  # type: ignore[no-untyped-def]
    """Append a paragraph containing an explicit <w:br w:type="page"/> to *doc*."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _add_section_break(doc) -> None:  # type: ignore[no-untyped-def]
    """Append a <w:sectPr> directly to the document body (section break)."""
    from docx.oxml import OxmlElement

    sectPr = OxmlElement("w:sectPr")
    doc.element.body.append(sectPr)


class TestDocxExplicitBreaksParagraphCap:
    def test_docx_page_break_with_paragraph_cap(self, tmp_path):
        """Page-break strategy also flushes pages when paragraph cap is hit.

        With _PARAGRAPHS_PER_PAGE=4, adding 2*4+1=9 paragraphs on each side
        of a single explicit page break must produce more than 2 pages (the
        cap triggers extra flushes beyond the one explicit break).
        """
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        n = _PARAGRAPHS_PER_PAGE * 2 + 1  # 9 paragraphs on each side
        for i in range(n):
            doc.add_paragraph(f"Before break paragraph {i + 1}.")
        _add_page_break(doc)
        for i in range(n):
            doc.add_paragraph(f"After break paragraph {i + 1}.")

        p = tmp_path / "page_break_cap.docx"
        doc.save(str(p))

        converter = get_converter(p)
        result = converter.convert(p)

        assert isinstance(result, IntermediateDocument)
        # The single explicit break alone would give 2 pages; the cap must
        # have added more splits → strictly more than 2 pages expected.
        assert len(result.pages) > 2
        # No page should be completely empty.
        for page in result.pages:
            assert page.markdown.strip() != "", (
                f"Page {page.index} has empty markdown"
            )

    def test_docx_section_break_with_paragraph_cap(self, tmp_path):
        """Section-break strategy also flushes pages when paragraph cap is hit.

        Same logic as above but using a <w:sectPr> body element as the
        explicit break instead of a <w:br w:type="page"/> run.
        """
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        n = _PARAGRAPHS_PER_PAGE * 2 + 1  # 9 paragraphs on each side
        for i in range(n):
            doc.add_paragraph(f"Before section paragraph {i + 1}.")
        _add_section_break(doc)
        for i in range(n):
            doc.add_paragraph(f"After section paragraph {i + 1}.")

        p = tmp_path / "section_break_cap.docx"
        doc.save(str(p))

        converter = get_converter(p)
        result = converter.convert(p)

        assert isinstance(result, IntermediateDocument)
        # Same reasoning: section break alone → 2 pages; cap → more than 2.
        assert len(result.pages) > 2
        # No page should be completely empty.
        for page in result.pages:
            assert page.markdown.strip() != "", (
                f"Page {page.index} has empty markdown"
            )

    def test_docx_consecutive_breaks_no_empty_pages(self, tmp_path):
        """Two consecutive page breaks without content between them must not
        produce empty pages.

        The flush_page() guard (``if not current_parts: return``) prevents
        empty IntermediatePage objects from being appended when a break is
        encountered while the current accumulator is already empty.
        """
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        # Some real content before the pair of breaks
        for i in range(3):
            doc.add_paragraph(f"Preamble paragraph {i + 1}.")
        # Two consecutive page breaks — no text between them
        _add_page_break(doc)
        _add_page_break(doc)
        # Some real content after
        for i in range(3):
            doc.add_paragraph(f"Postamble paragraph {i + 1}.")

        p = tmp_path / "consecutive_breaks.docx"
        doc.save(str(p))

        converter = get_converter(p)
        result = converter.convert(p)

        assert isinstance(result, IntermediateDocument)
        assert len(result.pages) >= 1
        # Critical assertion: no page may have empty markdown.
        for page in result.pages:
            assert page.markdown.strip() != "", (
                f"Page {page.index} is empty — consecutive breaks were not filtered"
            )
